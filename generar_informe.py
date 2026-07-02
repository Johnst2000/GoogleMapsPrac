from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

BASE = Path(__file__).parent
OUT = BASE / "Informe_Mapas_Lugares_Turisticos_Quevedo.docx"

FILES = {
    "MainActivity2.java": BASE / "app/src/main/java/com/uteq/software/googlemaps/MainActivity2.java",
    "activity_main2.xml": BASE / "app/src/main/res/layout/activity_main2.xml",
    "AndroidManifest.xml": BASE / "app/src/main/AndroidManifest.xml",
    "MainActivity.java": BASE / "app/src/main/java/com/uteq/software/googlemaps/MainActivity.java",
}


def add_title(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_normal(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_code_block(doc, code, filename):
    add_normal(doc, filename, bold=True)
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.space_after = Pt(12)


def add_screenshot_placeholder(doc, title, hint):
    add_title(doc, title, level=3)
    add_normal(doc, hint)
    p = doc.add_paragraph()
    run = p.add_run("[ PEGAR CAPTURA AQUÍ ]")
    run.bold = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def main():
    doc = Document()

    title = doc.add_heading("Mapa de Lugares Turísticos de Quevedo", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "Aplicaciones Móviles - 6to Semestre",
        "Estudiante: John Silva",
        "Correo: jsilvat2@uteq.edu.ec",
        "Docente: Cristian Zambrano",
        "Proyecto: GoogleMapsPrac",
    ]:
        info.add_run(line + "\n")

    doc.add_page_break()

    add_title(doc, "1. Capturas de la aplicación en funcionamiento", 1)
    add_normal(
        doc,
        "Pegar en cada espacio las capturas del emulador o dispositivo mientras la app está en ejecución.",
    )

    add_screenshot_placeholder(
        doc,
        "Captura 1 - Interfaz del mapa (3 pts)",
        "Mapa en vista satélite, círculo rojo, slider de radio, latitud y longitud actualizadas al mover el mapa.",
    )
    add_screenshot_placeholder(
        doc,
        "Captura 2 - Categorías y subcategorías (2 pts)",
        "Spinners con la lista cargada desde la API. Mostrar una categoría seleccionada y sus subcategorías.",
    )
    add_screenshot_placeholder(
        doc,
        "Captura 3 - Lugares turísticos en el mapa (3 pts)",
        "Marcadores dentro del círculo con nombres visibles al tocar un pin.",
    )
    add_screenshot_placeholder(
        doc,
        "Captura 4 - Filtros aplicados (2 pts)",
        "Ejemplo: Categoría Alimentación y Bebidas + Subcategoría ComidasRapidas mostrando solo esos lugares.",
    )

    doc.add_page_break()

    add_title(doc, "2. Código fuente de la actividad principal", 1)
    add_normal(doc, "Actividad: MainActivity2.java (Java)")
    add_code_block(doc, FILES["MainActivity2.java"].read_text(encoding="utf-8"), "MainActivity2.java")

    doc.add_page_break()

    add_title(doc, "3. Diseño en XML", 1)
    add_normal(doc, "Layout: activity_main2.xml")
    add_code_block(doc, FILES["activity_main2.xml"].read_text(encoding="utf-8"), "activity_main2.xml")

    doc.add_page_break()

    add_title(doc, "4. AndroidManifest.xml", 1)
    add_code_block(doc, FILES["AndroidManifest.xml"].read_text(encoding="utf-8"), "AndroidManifest.xml")

    doc.add_page_break()

    add_title(doc, "5. Otras clases del proyecto", 1)
    add_normal(
        doc,
        "MainActivity.java: práctica anterior del silabo (polilínea y marcador UTEQ, páginas 16-18).",
    )
    add_code_block(doc, FILES["MainActivity.java"].read_text(encoding="utf-8"), "MainActivity.java")

    doc.add_page_break()

    add_title(doc, "6. Rúbrica de evaluación", 1)

    rubrica = [
        ("Interfaz del Mapa: Círculos, Slider y Mapa", "3", [
            "SupportMapFragment con Google Maps",
            "Círculo rojo con radio en metros (slider)",
            "Latitud y longitud se actualizan al mover el mapa",
            "Vista satélite y controles de zoom",
        ]),
        ("Llenar lista de Categorías y Subcategorías", "2", [
            "Volley GET a categoria/getlistadoCB",
            "Volley GET a subcategoria/getlistadoCB/{id}",
            "Spinners Categoria y Subcategoria",
        ]),
        ("Obtener lugares turísticos y mostrarlos en el mapa", "3", [
            "Volley GET a lugar_turistico/json_getlistadoMapa",
            "Marcadores con nombre en el mapa",
            "Solo lugares dentro del círculo",
        ]),
        ("Aplicar filtros a la búsqueda", "2", [
            "Parámetros idc y idsc en la URL",
            "Actualización al cambiar categoría o subcategoría",
            "Actualización al mover mapa o cambiar radio",
        ]),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Criterio"
    hdr[1].text = "Puntos"
    hdr[2].text = "Implementado en el proyecto"

    for criterio, puntos, items in rubrica:
        row = table.add_row().cells
        row[0].text = criterio
        row[1].text = puntos
        row[2].text = "\n".join(f"• {i}" for i in items)

    doc.add_paragraph()
    add_title(doc, "7. Endpoints utilizados", 2)
    for url in [
        "http://35.153.103.86/turismo10022025/categoria/getlistadoCB",
        "http://35.153.103.86/turismo10022025/subcategoria/getlistadoCB/{idCategoria}",
        "http://35.153.103.86/turismo10022025/lugar_turistico/json_getlistadoMapa?lat=&lng=&radio=&idc=&idsc=",
    ]:
        doc.add_paragraph(url, style="List Bullet")

    add_title(doc, "8. Dependencias (build.gradle)", 2)
    deps = """implementation libs.play.services.maps
implementation libs.play.services.location
implementation libs.appcompat
implementation libs.constraintlayout
implementation libs.material
implementation libs.volley"""
    add_code_block(doc, deps, "app/build.gradle - dependencies")

    doc.save(OUT)
    print(f"Documento creado: {OUT}")


if __name__ == "__main__":
    main()
